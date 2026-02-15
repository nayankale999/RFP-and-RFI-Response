#!/usr/bin/env python3
"""
RFP/RFI Schedule Extractor

Parses an RFP/RFI document (PDF or DOCX) and uses Claude AI to extract
the procurement schedule — key dates, deadlines, and milestones.

Usage:
    python3 extract_schedule.py --input document.pdf [--output schedule.json] [--format json|csv|markdown] [--verbose]
"""

import argparse
import csv
import io
import json
import logging
import os
import sys
from datetime import datetime
from typing import Any, Optional

logger = logging.getLogger(__name__)

# Maximum characters to send to Claude (covers most schedule locations)
MAX_TEXT_CHARS = 15000


# ---------------------------------------------------------------------------
# Document Parser
# ---------------------------------------------------------------------------

class DocumentParser:
    """Parses PDF and DOCX documents to extract text and tables."""

    @staticmethod
    def parse(file_path: str) -> dict[str, Any]:
        """Parse a document and return text + tables.

        Returns:
            {
                "text": str,
                "tables": list[list[list[str]]],
                "filename": str,
                "page_count": int,
            }
        """
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return DocumentParser._parse_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return DocumentParser._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .docx")

    @staticmethod
    def _parse_pdf(file_path: str) -> dict[str, Any]:
        """Parse a PDF file using pdfplumber."""
        import pdfplumber

        all_text = []
        all_tables = []

        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                # Extract text
                text = page.extract_text()
                if text:
                    all_text.append(text)

                # Extract tables
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        # Clean up None values
                        cleaned = []
                        for row in table:
                            cleaned.append([str(cell).strip() if cell else "" for cell in row])
                        all_tables.append(cleaned)

        return {
            "text": "\n\n".join(all_text),
            "tables": all_tables,
            "filename": os.path.basename(file_path),
            "page_count": page_count,
        }

    @staticmethod
    def _parse_docx(file_path: str) -> dict[str, Any]:
        """Parse a DOCX file using python-docx."""
        from docx import Document

        doc = Document(file_path)
        all_text = []
        all_tables = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text.strip())

        # Extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            all_tables.append(rows)

        return {
            "text": "\n".join(all_text),
            "tables": all_tables,
            "filename": os.path.basename(file_path),
            "page_count": 0,  # DOCX doesn't have page count easily
        }


# ---------------------------------------------------------------------------
# Schedule Extractor (Claude AI)
# ---------------------------------------------------------------------------

EXTRACTION_TOOL = {
    "name": "extract_schedule",
    "description": "Extract procurement schedule events and dates from an RFP/RFI document",
    "input_schema": {
        "type": "object",
        "properties": {
            "schedule_events": {
                "type": "array",
                "description": "List of schedule events extracted from the document",
                "items": {
                    "type": "object",
                    "properties": {
                        "event_type": {
                            "type": "string",
                            "enum": [
                                "rfp_release",
                                "intention_to_respond",
                                "clarification_deadline",
                                "submission_deadline",
                                "evaluation",
                                "shortlist_notification",
                                "demo_presentation",
                                "poc_start",
                                "poc_end",
                                "selection_decision",
                                "contracting",
                                "implementation_start",
                                "other",
                            ],
                            "description": "Category of the schedule event",
                        },
                        "event_name": {
                            "type": "string",
                            "description": "Human-readable name of the event as stated in the document",
                        },
                        "date": {
                            "type": "string",
                            "description": "Date in YYYY-MM-DD format for exact dates, or descriptive for week numbers/relative dates (e.g., 'Week 49', 'Within 30 days of submission')",
                        },
                        "date_type": {
                            "type": "string",
                            "enum": ["exact", "approximate", "week_number", "relative", "tbd"],
                            "description": "Whether the date is exact, approximate, a week number, relative, or to be determined",
                        },
                        "is_deadline": {
                            "type": "boolean",
                            "description": "True if this is a deadline (must be completed BY this date), false if it's a start date",
                        },
                        "notes": {
                            "type": "string",
                            "description": "Additional context, conditions, or instructions related to this event",
                        },
                    },
                    "required": ["event_type", "event_name", "date", "date_type", "is_deadline"],
                },
            },
            "source_section": {
                "type": "string",
                "description": "The section or heading where the schedule was found in the document (e.g., 'Section 4.4 Planning', 'Schedule and Timeline')",
            },
            "additional_notes": {
                "type": "string",
                "description": "Any caveats or notes about the schedule (e.g., 'timetable is indicative', 'dates subject to change')",
            },
        },
        "required": ["schedule_events", "source_section"],
    },
}

SYSTEM_PROMPT = """You are an expert at extracting procurement schedules from RFP (Request for Proposal) and RFI (Request for Information) documents.

Your task is to extract ALL dates and timeline events related to the procurement process. Look carefully for:

1. RFP/RFI issuance or release date
2. Deadline to confirm intention to respond (Yes/No reply)
3. Questions/clarifications submission deadline
4. Answers to questions deadline
5. Proposal/response submission deadline
6. Evaluation period or dates
7. Shortlist or initial selection notification
8. Vendor demos, presentations, or site visits
9. Proof of Concept (PoC) start and end dates
10. Final selection/award decision date
11. Contracting/negotiation period
12. Implementation or project start date
13. Any other milestone dates mentioned

The schedule information may appear in:
- A formal table (most common)
- Bullet points or numbered lists
- Narrative paragraphs
- Multiple locations throughout the document

Extract from ALL of these sources. Do not miss any dates.

DATE FORMAT RULES:
- For exact dates (e.g., "October 16th, 2024"): Convert to YYYY-MM-DD format (2024-10-16)
- For week numbers (e.g., "Week 49"): Keep as-is (e.g., "Week 49")
- For relative dates (e.g., "within 30 days"): Keep the relative description
- For approximate dates (e.g., "mid-November"): Provide your best estimate in YYYY-MM-DD and mark date_type as "approximate"

DEADLINE DETECTION:
- Words like "before", "by", "no later than", "due", "deadline" indicate a deadline (is_deadline = true)
- Words like "start", "begin", "commence", "from" indicate a start date (is_deadline = false)
- If a table has "Start" and "Before" columns, the "Before" column entries are deadlines

Always note any caveats mentioned about the schedule (e.g., "indicative", "subject to change", "right to modify")."""


class ScheduleExtractor:
    """Extracts schedule events from parsed document content using Claude AI."""

    def __init__(self, api_key: str, model: str = "claude-sonnet-4-5-20250929"):
        import anthropic
        self.client = anthropic.Anthropic(api_key=api_key)
        self.model = model

    def extract(self, parsed_doc: dict[str, Any]) -> dict[str, Any]:
        """Extract schedule from parsed document content.

        Args:
            parsed_doc: Output from DocumentParser.parse()

        Returns:
            Structured schedule data with events, source, and notes.
        """
        # Build the user prompt with both text and tables
        user_content = self._build_prompt(parsed_doc)

        logger.info(f"Sending {len(user_content)} chars to Claude for extraction...")

        response = self.client.messages.create(
            model=self.model,
            max_tokens=4096,
            system=SYSTEM_PROMPT,
            tools=[EXTRACTION_TOOL],
            tool_choice={"type": "tool", "name": "extract_schedule"},
            messages=[
                {"role": "user", "content": user_content},
            ],
        )

        # Extract structured data from tool use response
        for block in response.content:
            if block.type == "tool_use" and block.name == "extract_schedule":
                result = block.input
                # Add metadata
                result["document"] = parsed_doc.get("filename", "unknown")
                result["extracted_at"] = datetime.now().isoformat()
                return result

        raise RuntimeError("Claude did not return a tool_use response")

    def _build_prompt(self, parsed_doc: dict[str, Any]) -> str:
        """Build the user prompt combining text and tables."""
        parts = []

        parts.append(
            f"Extract all schedule events and procurement dates from this "
            f"RFP/RFI document ({parsed_doc.get('filename', 'unknown')})."
        )
        parts.append("")

        # Include tables first (most likely to contain schedule)
        tables = parsed_doc.get("tables", [])
        if tables:
            parts.append("=== TABLES FOUND IN DOCUMENT ===")
            for i, table in enumerate(tables, 1):
                parts.append(f"\nTable {i}:")
                for row in table:
                    parts.append(" | ".join(row))
            parts.append("")

        # Include document text (truncated if needed)
        text = parsed_doc.get("text", "")
        if len(text) > MAX_TEXT_CHARS:
            parts.append(f"=== DOCUMENT TEXT (first {MAX_TEXT_CHARS} characters) ===")
            parts.append(text[:MAX_TEXT_CHARS])
            parts.append("\n[... text truncated ...]")
        else:
            parts.append("=== FULL DOCUMENT TEXT ===")
            parts.append(text)

        parts.append("")
        parts.append(
            "Extract every schedule event, date, and deadline from the above content. "
            "Do not miss any dates mentioned anywhere in the document."
        )

        return "\n".join(parts)


# ---------------------------------------------------------------------------
# Output Formatters
# ---------------------------------------------------------------------------

def format_markdown_table(result: dict[str, Any]) -> str:
    """Format extraction result as a markdown table."""
    lines = []

    lines.append(f"**Schedule extracted from:** {result.get('document', 'unknown')}")
    lines.append(f"**Source:** {result.get('source_section', 'N/A')}")
    lines.append("")

    events = result.get("schedule_events", [])
    if not events:
        lines.append("No schedule events found.")
        return "\n".join(lines)

    # Header
    lines.append("| # | Event | Date | Type | Deadline? | Notes |")
    lines.append("|---|-------|------|------|-----------|-------|")

    for i, event in enumerate(events, 1):
        name = event.get("event_name", "N/A")
        date = event.get("date", "TBD")
        date_type = event.get("date_type", "")
        is_deadline = "Yes" if event.get("is_deadline") else "No"
        notes = event.get("notes", "")
        # Truncate long notes for table display
        if len(notes) > 60:
            notes = notes[:57] + "..."
        lines.append(f"| {i} | {name} | {date} | {date_type} | {is_deadline} | {notes} |")

    # Additional notes
    additional = result.get("additional_notes", "")
    if additional:
        lines.append("")
        lines.append(f"**Note:** {additional}")

    return "\n".join(lines)


def format_csv(result: dict[str, Any]) -> str:
    """Format extraction result as CSV."""
    output = io.StringIO()
    writer = csv.writer(output)

    # Header
    writer.writerow(["#", "Event Type", "Event Name", "Date", "Date Type", "Is Deadline", "Notes"])

    events = result.get("schedule_events", [])
    for i, event in enumerate(events, 1):
        writer.writerow([
            i,
            event.get("event_type", ""),
            event.get("event_name", ""),
            event.get("date", ""),
            event.get("date_type", ""),
            "Yes" if event.get("is_deadline") else "No",
            event.get("notes", ""),
        ])

    return output.getvalue()


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Extract procurement schedule from RFP/RFI documents (PDF or DOCX)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
    python3 extract_schedule.py --input rfp.pdf
    python3 extract_schedule.py --input rfp.pdf --output schedule.json --format json
    python3 extract_schedule.py --input proposal.docx --format csv --output schedule.csv
        """,
    )
    parser.add_argument("--input", required=True, help="Path to PDF or DOCX document")
    parser.add_argument("--output", help="Output file path (optional, prints to stdout if omitted)")
    parser.add_argument(
        "--format",
        choices=["json", "csv", "markdown"],
        default="markdown",
        help="Output format (default: markdown)",
    )
    parser.add_argument("--model", default="claude-sonnet-4-5-20250929", help="Claude model to use")
    parser.add_argument(
        "--parse-only",
        action="store_true",
        help="Only parse the document and output text + tables as JSON (no AI extraction). "
             "Use this when Claude Code will handle the AI extraction directly.",
    )
    parser.add_argument("--verbose", action="store_true", help="Enable verbose logging")
    args = parser.parse_args()

    # Configure logging
    log_level = logging.DEBUG if args.verbose else logging.WARNING
    logging.basicConfig(
        level=log_level,
        format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    )

    # Step 1: Parse document
    print(f"Parsing document: {args.input}", file=sys.stderr)
    try:
        parsed = DocumentParser.parse(args.input)
    except Exception as e:
        print(f"Error parsing document: {e}", file=sys.stderr)
        sys.exit(1)

    page_info = f" ({parsed['page_count']} pages)" if parsed.get("page_count") else ""
    table_info = f", {len(parsed['tables'])} tables" if parsed.get("tables") else ""
    print(f"  Parsed: {len(parsed['text'])} chars{page_info}{table_info}", file=sys.stderr)

    # --parse-only mode: output raw parsed content for Claude to analyse directly
    if args.parse_only:
        # Build structured output with text and tables
        output = {
            "filename": parsed.get("filename", "unknown"),
            "page_count": parsed.get("page_count", 0),
            "table_count": len(parsed.get("tables", [])),
        }

        # Format tables as readable text
        tables = parsed.get("tables", [])
        if tables:
            table_text_parts = []
            for i, table in enumerate(tables, 1):
                table_text_parts.append(f"=== Table {i} ===")
                for row in table:
                    table_text_parts.append(" | ".join(row))
            output["tables_text"] = "\n".join(table_text_parts)

        # Include full text (truncated for very large docs)
        text = parsed.get("text", "")
        if len(text) > MAX_TEXT_CHARS:
            output["document_text"] = text[:MAX_TEXT_CHARS]
            output["text_truncated"] = True
        else:
            output["document_text"] = text
            output["text_truncated"] = False

        print(json.dumps(output, indent=2, ensure_ascii=False))
        return

    # Full extraction mode: requires API key
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("Error: ANTHROPIC_API_KEY environment variable is not set.", file=sys.stderr)
        print("Set it with: export ANTHROPIC_API_KEY='your-key-here'", file=sys.stderr)
        print("Alternatively, use --parse-only to just extract text/tables and let Claude Code handle analysis.", file=sys.stderr)
        sys.exit(1)

    # Step 2: Extract schedule via Claude
    print("Extracting schedule using Claude AI...", file=sys.stderr)
    extractor = ScheduleExtractor(api_key=api_key, model=args.model)
    try:
        result = extractor.extract(parsed)
    except Exception as e:
        print(f"Error during extraction: {e}", file=sys.stderr)
        if args.verbose:
            import traceback
            traceback.print_exc()
        sys.exit(1)

    event_count = len(result.get("schedule_events", []))
    print(f"  Found {event_count} schedule events", file=sys.stderr)

    # Step 3: Format output
    if args.format == "json":
        formatted = json.dumps(result, indent=2, ensure_ascii=False)
    elif args.format == "csv":
        formatted = format_csv(result)
    else:  # markdown
        formatted = format_markdown_table(result)

    # Step 4: Output
    if args.output:
        os.makedirs(os.path.dirname(os.path.abspath(args.output)), exist_ok=True)
        with open(args.output, "w", encoding="utf-8") as f:
            f.write(formatted)
        print(f"Output saved to: {args.output}", file=sys.stderr)
    else:
        print(formatted)


if __name__ == "__main__":
    main()
