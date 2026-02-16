#!/usr/bin/env python3
"""
RFP Win Plan DOCX Generator

Generates a professional Win Plan Word document from extracted schedule data.
The Win Plan serves as the internal team's action plan for responding to an RFP/RFI.

Usage:
    python3 generate_win_plan.py --input schedule_data.json --output Win_Plan.docx
"""

import argparse
import json
import logging
import os
import sys
from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)

# Colors matching the corporate branding
HEADING_BAR_COLOR = "314662"  # Dark slate blue
TABLE_HEADER_BG = "D9D9D9"   # Light gray
DEADLINE_HIGHLIGHT = "FFF3CD" # Light yellow for deadlines
WHITE = "FFFFFF"
LIGHT_BLUE_BG = "EBF0F5"     # Very light blue for alternating rows


def set_cell_shading(cell, color: str):
    """Set background shading color on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_bar(doc: Document, text: str):
    """Add a dark blue heading bar with white text (matching CreateRFIResponse style)."""
    # Add a single-cell table to simulate the heading bar
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""

    # Set cell background
    set_cell_shading(cell, HEADING_BAR_COLOR)

    # Set cell width to full page
    cell.width = Cm(17)

    # Add text with white formatting
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(14)
    run.font.name = "Calibri"
    run.bold = True

    # Reduce spacing
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)

    # Remove table borders (keep only the fill)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    doc.add_paragraph("")  # spacer


def style_table_header_row(table, col_count: int):
    """Style the first row of a table as a gray header."""
    for i in range(col_count):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, TABLE_HEADER_BG)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = "Calibri"


class WinPlanGenerator:
    """Generates a professional RFP Win Plan DOCX document."""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Configure document styles."""
        # Normal style
        style = self.doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)

    def generate(self, data: dict[str, Any], output_path: str) -> str:
        """Generate the complete Win Plan document."""
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

        self._add_cover(data)
        self._add_rfp_overview(data)
        self._add_procurement_schedule(data)
        self._add_key_deadlines(data)
        self._add_response_team()
        self._add_win_strategy(data)
        self._add_action_items()
        self._add_notes()
        self._add_footer(data)

        self.doc.save(output_path)
        return output_path

    def _add_cover(self, data: dict):
        """Section: Cover / Title page."""
        # Add some spacing
        for _ in range(4):
            self.doc.add_paragraph("")

        # Title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("RFP WIN PLAN")
        run.font.size = Pt(32)
        run.font.color.rgb = RGBColor(0x31, 0x46, 0x62)
        run.bold = True
        run.font.name = "Calibri"

        self.doc.add_paragraph("")

        # Client name
        client = self.doc.add_paragraph()
        client.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = client.add_run(data.get("client_name", "Client"))
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x31, 0x46, 0x62)
        run.font.name = "Calibri"

        # RFP title
        rfp_title = self.doc.add_paragraph()
        rfp_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = rfp_title.add_run(data.get("rfp_title", "RFP Response"))
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.font.name = "Calibri"

        self.doc.add_paragraph("")

        # Date
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"Date: {datetime.now().strftime('%d %B %Y')}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

        # Version
        version = self.doc.add_paragraph()
        version.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = version.add_run("Version 1.0 | CONFIDENTIAL")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        self.doc.add_page_break()

    def _add_rfp_overview(self, data: dict):
        """Section: RFP Overview."""
        add_heading_bar(self.doc, "RFP Overview")

        # Overview details table
        details = [
            ("Client", data.get("client_name", "—")),
            ("RFP Title", data.get("rfp_title", "—")),
            ("Source Document", data.get("document", "—")),
            ("Pages", str(data.get("page_count", "—"))),
            ("Schedule Source", data.get("source_section", "—")),
        ]

        table = self.doc.add_table(rows=len(details), cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, (label, value) in enumerate(details):
            # Label cell
            cell_label = table.rows[i].cells[0]
            cell_label.text = ""
            set_cell_shading(cell_label, LIGHT_BLUE_BG)
            run = cell_label.paragraphs[0].add_run(label)
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = "Calibri"

            # Value cell
            cell_value = table.rows[i].cells[1]
            cell_value.text = ""
            run = cell_value.paragraphs[0].add_run(value)
            run.font.size = Pt(10)
            run.font.name = "Calibri"

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Cm(5)
            row.cells[1].width = Cm(12)

        # Additional notes
        notes = data.get("additional_notes", "")
        if notes:
            self.doc.add_paragraph("")
            p = self.doc.add_paragraph()
            run = p.add_run("Note: ")
            run.bold = True
            run.font.size = Pt(10)
            run = p.add_run(notes)
            run.font.size = Pt(10)
            run.italic = True

        self.doc.add_paragraph("")

    def _add_procurement_schedule(self, data: dict):
        """Section: Full Procurement Schedule table."""
        add_heading_bar(self.doc, "Procurement Schedule")

        events = data.get("events", data.get("schedule_events", []))
        if not events:
            self.doc.add_paragraph("No schedule events extracted.")
            return

        # Create table
        headers = ["#", "Event", "Date", "Type", "Deadline?", "Notes"]
        table = self.doc.add_table(rows=1 + len(events), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for j, header in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(header)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = "Calibri"

        style_table_header_row(table, len(headers))

        # Data rows
        for i, event in enumerate(events):
            row = table.rows[i + 1]
            values = [
                str(i + 1),
                event.get("event_name", "—"),
                event.get("date", "TBD"),
                event.get("date_type", "—"),
                "Yes" if event.get("is_deadline") else "No",
                event.get("notes", ""),
            ]
            for j, val in enumerate(values):
                cell = row.cells[j]
                cell.text = ""
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(9)
                run.font.name = "Calibri"

            # Highlight deadline rows
            if event.get("is_deadline"):
                for j in range(len(headers)):
                    set_cell_shading(row.cells[j], DEADLINE_HIGHLIGHT)

        # Set column widths
        col_widths = [Cm(1), Cm(5.5), Cm(3), Cm(2), Cm(1.5), Cm(4)]
        for row in table.rows:
            for j, width in enumerate(col_widths):
                row.cells[j].width = width

        self.doc.add_paragraph("")

    def _add_key_deadlines(self, data: dict):
        """Section: Key Deadlines Summary — filtered to deadlines only."""
        add_heading_bar(self.doc, "Key Deadlines Summary")

        events = data.get("events", data.get("schedule_events", []))
        deadlines = [e for e in events if e.get("is_deadline")]

        if not deadlines:
            self.doc.add_paragraph("No specific deadlines identified.")
            return

        p = self.doc.add_paragraph()
        run = p.add_run(f"{len(deadlines)} key deadlines identified. ")
        run.font.size = Pt(10)
        run = p.add_run("Ensure the response team is aware of these critical dates.")
        run.font.size = Pt(10)
        run.italic = True

        self.doc.add_paragraph("")

        # Deadlines table (simplified)
        table = self.doc.add_table(rows=1 + len(deadlines), cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["Priority", "Deadline", "Date", "Action Required"]
        for j, header in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(header)
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = "Calibri"

        style_table_header_row(table, 4)

        for i, event in enumerate(deadlines):
            row = table.rows[i + 1]
            # Priority based on event type
            priority_map = {
                "submission_deadline": "CRITICAL",
                "intention_to_respond": "HIGH",
                "clarification_deadline": "HIGH",
                "poc_end": "HIGH",
                "shortlist_notification": "MEDIUM",
                "selection_decision": "MEDIUM",
                "contracting": "MEDIUM",
                "implementation_start": "LOW",
            }
            priority = priority_map.get(event.get("event_type", ""), "MEDIUM")

            values = [
                priority,
                event.get("event_name", "—"),
                event.get("date", "TBD"),
                event.get("notes", "Review and prepare"),
            ]

            for j, val in enumerate(values):
                cell = row.cells[j]
                cell.text = ""
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(9)
                run.font.name = "Calibri"

                # Color-code priority
                if j == 0:
                    run.font.bold = True
                    if val == "CRITICAL":
                        run.font.color.rgb = RGBColor(0xDC, 0x35, 0x45)  # Red
                    elif val == "HIGH":
                        run.font.color.rgb = RGBColor(0xFD, 0x7E, 0x14)  # Orange
                    elif val == "MEDIUM":
                        run.font.color.rgb = RGBColor(0xFF, 0xC1, 0x07)  # Amber

        col_widths = [Cm(2.5), Cm(5.5), Cm(3), Cm(6)]
        for row in table.rows:
            for j, width in enumerate(col_widths):
                row.cells[j].width = width

        self.doc.add_paragraph("")

    def _add_response_team(self):
        """Section: Response Team (placeholder)."""
        add_heading_bar(self.doc, "Response Team")

        p = self.doc.add_paragraph()
        run = p.add_run("Assign team members to each role below. Update as the response progresses.")
        run.font.size = Pt(10)
        run.italic = True

        self.doc.add_paragraph("")

        roles = [
            ("Bid Manager / Proposal Lead", "", "Overall response coordination, timeline management"),
            ("Solution Architect", "", "Technical solution design, architecture documentation"),
            ("Pre-Sales / Demo Lead", "", "Solution demonstrations, PoC execution"),
            ("Subject Matter Expert (SME)", "", "Domain expertise, functional responses"),
            ("Commercial / Pricing Lead", "", "Pricing model, TCO calculation, commercial terms"),
            ("Legal", "", "Contract review, T&Cs, compliance checks"),
            ("Executive Sponsor", "", "Strategic oversight, escalation point, executive summary"),
        ]

        table = self.doc.add_table(rows=1 + len(roles), cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["Role", "Name", "Responsibility"]
        for j, header in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(header)
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = "Calibri"

        style_table_header_row(table, 3)

        for i, (role, name, resp) in enumerate(roles):
            row = table.rows[i + 1]
            for j, val in enumerate([role, name, resp]):
                cell = row.cells[j]
                cell.text = ""
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(9)
                run.font.name = "Calibri"
                if j == 1 and not val:
                    # Placeholder for name
                    run = cell.paragraphs[0].add_run("[To be assigned]")
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                    run.italic = True

        col_widths = [Cm(5), Cm(4), Cm(8)]
        for row in table.rows:
            for j, width in enumerate(col_widths):
                row.cells[j].width = width

        self.doc.add_paragraph("")

    def _add_win_strategy(self, data: dict):
        """Section: Win Strategy — pre-filled with solution data when available."""
        add_heading_bar(self.doc, "Win Strategy")

        # Solution overview paragraph (if provided)
        solution_overview = data.get("solution_overview", "")
        if solution_overview:
            p = self.doc.add_paragraph()
            run = p.add_run("Solution: ")
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            run = p.add_run(data.get("solution_name", ""))
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            run.bold = True

            p = self.doc.add_paragraph()
            run = p.add_run(solution_overview)
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            self.doc.add_paragraph("")

        # Each strategy section: render bullet list if data exists, else placeholder
        sections = [
            (
                "Key Differentiators",
                "differentiators",
                "What makes our solution uniquely suited for this client? List 3-5 differentiators.",
            ),
            (
                "Competitive Advantages",
                "competitive_advantages",
                "How do we compare against likely competitors? What are our strengths?",
            ),
            (
                "Client Pain Points",
                None,
                "What are the client's primary challenges? How does our solution address each one?",
            ),
            (
                "Risk Areas",
                "risk_areas",
                "What could weaken our proposal? Pricing, experience gaps, technical gaps?",
            ),
            (
                "Win Themes",
                "win_themes",
                "What 2-3 key messages should run throughout our response?",
            ),
        ]

        for title, data_key, fallback_prompt in sections:
            # Section title
            p = self.doc.add_paragraph()
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x31, 0x46, 0x62)

            items = data.get(data_key, []) if data_key else []
            if items and isinstance(items, list):
                # Render as bullet points
                for item in items:
                    bp = self.doc.add_paragraph(style="List Bullet")
                    run = bp.add_run(str(item))
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
            else:
                # Fallback: show placeholder prompt
                p = self.doc.add_paragraph()
                run = p.add_run(fallback_prompt)
                run.font.size = Pt(10)
                run.italic = True
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            self.doc.add_paragraph("")

        self.doc.add_paragraph("")

    def _add_action_items(self):
        """Section: Action Items (placeholder)."""
        add_heading_bar(self.doc, "Action Items")

        p = self.doc.add_paragraph()
        run = p.add_run("Track all actions required to complete the RFP response.")
        run.font.size = Pt(10)
        run.italic = True

        self.doc.add_paragraph("")

        table = self.doc.add_table(rows=6, cols=5)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["#", "Action", "Owner", "Due Date", "Status"]
        for j, header in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(header)
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = "Calibri"

        style_table_header_row(table, 5)

        # Pre-fill with common actions
        actions = [
            ("1", "Confirm intention to respond", "", "", "Not Started"),
            ("2", "Prepare demo / presentation", "", "", "Not Started"),
            ("3", "Draft response document", "", "", "Not Started"),
            ("4", "Review and finalise pricing", "", "", "Not Started"),
            ("5", "Submit final response", "", "", "Not Started"),
        ]

        for i, (num, action, owner, due, status) in enumerate(actions):
            row = table.rows[i + 1]
            for j, val in enumerate([num, action, owner, due, status]):
                cell = row.cells[j]
                cell.text = ""
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(9)
                run.font.name = "Calibri"

        col_widths = [Cm(1), Cm(6.5), Cm(3.5), Cm(3), Cm(3)]
        for row in table.rows:
            for j, width in enumerate(col_widths):
                row.cells[j].width = width

        self.doc.add_paragraph("")

    def _add_notes(self):
        """Section: Notes (empty)."""
        add_heading_bar(self.doc, "Notes")

        p = self.doc.add_paragraph()
        run = p.add_run("[Add any additional notes, observations, or meeting minutes here]")
        run.font.size = Pt(10)
        run.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Add some space for notes
        for _ in range(5):
            self.doc.add_paragraph("")

    def _add_footer(self, data: dict):
        """Add document footer."""
        section = self.doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run(
            f"CONFIDENTIAL | RFP Win Plan — {data.get('client_name', 'Client')} | "
            f"{datetime.now().strftime('%Y-%m-%d')}"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Generate an RFP Win Plan DOCX document from extracted schedule data",
    )
    parser.add_argument("--input", required=True, help="Path to schedule JSON data file")
    parser.add_argument("--output", required=True, help="Output DOCX file path")
    parser.add_argument("--verbose", action="store_true", help="Enable verbose logging")
    args = parser.parse_args()

    log_level = logging.DEBUG if args.verbose else logging.INFO
    logging.basicConfig(level=log_level, format="%(asctime)s [%(levelname)s] %(message)s")

    # Load data
    try:
        with open(args.input, "r", encoding="utf-8") as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f"Error: Input file not found: {args.input}", file=sys.stderr)
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"Error: Invalid JSON: {e}", file=sys.stderr)
        sys.exit(1)

    # Generate
    print(f"Generating Win Plan: {args.output}")
    generator = WinPlanGenerator()
    output_path = generator.generate(data, args.output)

    event_count = len(data.get("events", data.get("schedule_events", [])))
    deadline_count = len([e for e in data.get("events", data.get("schedule_events", [])) if e.get("is_deadline")])
    print(f"  Win Plan generated successfully!")
    print(f"  Events: {event_count} ({deadline_count} deadlines)")
    print(f"  Output: {output_path}")


if __name__ == "__main__":
    main()
