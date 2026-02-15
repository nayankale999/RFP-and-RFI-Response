import io
import logging

from docx import Document

from app.documents.parsers.base import BaseParser, ParsedDocument

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """Microsoft Word (.docx) parser."""

    def supported_extensions(self) -> list[str]:
        return [".docx", ".doc"]

    def parse(self, file_data: bytes, filename: str) -> ParsedDocument:
        try:
            doc = Document(io.BytesIO(file_data))
        except Exception as e:
            logger.error(f"Failed to parse DOCX {filename}: {e}")
            raise

        paragraphs = []
        sections = []
        current_section = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect headings
            if para.style and para.style.name.startswith("Heading"):
                level = int(para.style.name.replace("Heading ", "").replace("Heading", "1") or "1")
                if current_section:
                    sections.append(current_section)
                current_section = {"heading": text, "level": level, "content": ""}
            else:
                if current_section:
                    current_section["content"] += text + "\n"

            paragraphs.append(text)

        if current_section:
            sections.append(current_section)

        # Extract tables
        all_tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            all_tables.append(rows)

        full_text = "\n\n".join(paragraphs)
        return ParsedDocument(
            text=full_text,
            page_count=len(doc.sections),
            tables=all_tables,
            sections=sections,
            metadata={"filename": filename, "parser": "docx", "paragraph_count": len(paragraphs)},
        )
