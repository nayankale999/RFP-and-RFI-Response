import io
import logging
from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from app.shared.exceptions import ExportError

logger = logging.getLogger(__name__)


class WordGenerator:
    """Generates professional RFP response documents in Word format."""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Configure document styles for professional formatting."""
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # Configure heading styles
        for level in range(1, 4):
            heading_style = self.doc.styles[f"Heading {level}"]
            heading_style.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)  # Dark blue

    def generate(self, context: dict[str, Any]) -> bytes:
        """Generate the complete RFP response document.

        Args:
            context: Dict containing all response data:
                - project: Project metadata
                - company: Company information
                - requirements: List of requirements
                - responses: List of responses (keyed by requirement_id)
                - schedule: Schedule events
                - pricing: Pricing items
                - compliance_scores: Compliance scores
                - response_plan: Response plan

        Returns:
            Word document as bytes.
        """
        try:
            self._add_cover_page(context)
            self._add_table_of_contents()
            self._add_executive_summary(context)
            self._add_about_company(context)
            self._add_understanding_of_requirements(context)
            self._add_proposed_solution(context)
            self._add_functional_compliance_matrix(context)
            self._add_nfr_compliance_matrix(context)
            self._add_architecture_overview(context)
            self._add_implementation_approach(context)
            self._add_project_plan(context)
            self._add_pricing(context)
            self._add_assumptions(context)
            self._add_risks_mitigation(context)
            self._add_legal_compliance(context)
            self._add_footer(context)

            # Write to bytes
            buffer = io.BytesIO()
            self.doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()

        except Exception as e:
            logger.error(f"Word document generation failed: {e}")
            raise ExportError(f"Document generation failed: {e}")

    def _add_cover_page(self, ctx: dict):
        """Add branded cover page."""
        self.doc.add_paragraph("")
        self.doc.add_paragraph("")
        self.doc.add_paragraph("")

        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(ctx.get("project", {}).get("name", "RFP Response"))
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        run.bold = True

        self.doc.add_paragraph("")

        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"Prepared for: {ctx.get('project', {}).get('client_name', 'Client')}")
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        self.doc.add_paragraph("")

        company = ctx.get("company", {})
        details = self.doc.add_paragraph()
        details.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = details.add_run(f"Submitted by: {company.get('name', 'Company Name')}")
        run.font.size = Pt(14)

        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

        version = self.doc.add_paragraph()
        version.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = version.add_run("Version 1.0 | CONFIDENTIAL")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        self.doc.add_page_break()

    def _add_table_of_contents(self):
        """Add table of contents placeholder."""
        self.doc.add_heading("Table of Contents", level=1)
        p = self.doc.add_paragraph(
            "[ Table of Contents - Update this field after generating the document "
            "(Right-click > Update Field in Word) ]"
        )
        p.style.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        self.doc.add_page_break()

    def _add_executive_summary(self, ctx: dict):
        """Section 1: Executive Summary."""
        self.doc.add_heading("1. Executive Summary", level=1)

        project = ctx.get("project", {})
        scores = ctx.get("compliance_scores", {})

        self.doc.add_paragraph(
            f"This document presents our comprehensive response to the "
            f"{project.get('name', 'RFP')} issued by {project.get('client_name', 'the client')}. "
            f"We have carefully analyzed all requirements and prepared detailed responses "
            f"demonstrating our solution capabilities."
        )

        if scores:
            self.doc.add_paragraph(
                f"Our overall compliance score is {scores.get('overall_score', 'N/A')}%, "
                f"with {scores.get('status_breakdown', {}).get('fully_compliant', 0)} requirements "
                f"fully addressed out of {scores.get('total_requirements', 0)} total."
            )

        self.doc.add_paragraph("")

    def _add_about_company(self, ctx: dict):
        """Section 2: About the Company."""
        self.doc.add_heading("2. About the Company", level=1)
        company = ctx.get("company", {})
        self.doc.add_paragraph(
            company.get(
                "description",
                "[ Company description to be provided by the proposal team. Include company history, "
                "key differentiators, relevant experience, certifications, and industry expertise. ]",
            )
        )
        self.doc.add_paragraph("")

    def _add_understanding_of_requirements(self, ctx: dict):
        """Section 3: Understanding of Requirements."""
        self.doc.add_heading("3. Understanding of Requirements", level=1)

        requirements = ctx.get("requirements", [])
        type_counts = {}
        for req in requirements:
            t = req.get("type", "other")
            type_counts[t] = type_counts.get(t, 0) + 1

        self.doc.add_paragraph(
            f"We have identified and analyzed {len(requirements)} requirements across "
            f"the following categories:"
        )

        table = self.doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text = "Requirement Type"
        hdr[1].text = "Count"
        hdr[2].text = "Mandatory"

        for req_type, count in sorted(type_counts.items()):
            mandatory_count = sum(
                1 for r in requirements if r.get("type") == req_type and r.get("is_mandatory")
            )
            row = table.add_row().cells
            row[0].text = req_type.replace("_", " ").title()
            row[1].text = str(count)
            row[2].text = str(mandatory_count)

        self.doc.add_paragraph("")

    def _add_proposed_solution(self, ctx: dict):
        """Section 4: Proposed Solution Overview."""
        self.doc.add_heading("4. Proposed Solution Overview", level=1)
        self.doc.add_paragraph(
            "[ Proposed solution overview to be completed by the solution architecture team. "
            "Include high-level solution description, key components, technology stack, "
            "and differentiating features. ]"
        )
        self.doc.add_paragraph("")

    def _add_functional_compliance_matrix(self, ctx: dict):
        """Section 5: Functional Compliance Matrix."""
        self.doc.add_heading("5. Functional Compliance Matrix", level=1)
        self._add_compliance_table(ctx, "functional")

    def _add_nfr_compliance_matrix(self, ctx: dict):
        """Section 6: Non-Functional Compliance Matrix."""
        self.doc.add_heading("6. Non-Functional Compliance Matrix", level=1)
        self._add_compliance_table(ctx, "non_functional")

    def _add_compliance_table(self, ctx: dict, req_type: str):
        """Add a compliance matrix table for a requirement type."""
        requirements = [r for r in ctx.get("requirements", []) if r.get("type") == req_type]
        responses_by_id = {str(r.get("requirement_id")): r for r in ctx.get("responses", [])}

        if not requirements:
            self.doc.add_paragraph(f"No {req_type.replace('_', ' ')} requirements identified.")
            return

        table = self.doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "ID"
        hdr[1].text = "Requirement"
        hdr[2].text = "Category"
        hdr[3].text = "Compliance"
        hdr[4].text = "Response"

        for req in requirements:
            req_id = str(req.get("id", ""))
            resp = responses_by_id.get(req_id, {})
            row = table.add_row().cells
            row[0].text = req.get("req_number", "N/A")
            row[1].text = req.get("title", "N/A")[:80]
            row[2].text = (req.get("category") or "General").title()
            row[3].text = (resp.get("compliance_status") or "Pending").replace("_", " ").title()
            row[4].text = (resp.get("response_text") or "Response pending")[:200]

        self.doc.add_paragraph("")

    def _add_architecture_overview(self, ctx: dict):
        """Section 7: Architecture Overview."""
        self.doc.add_heading("7. Architecture Overview", level=1)
        self.doc.add_paragraph(
            "[ Architecture diagram and description to be provided by the architecture team. "
            "Include system architecture, deployment model, integration points, "
            "and security architecture. ]"
        )
        self.doc.add_paragraph("")

    def _add_implementation_approach(self, ctx: dict):
        """Section 8: Implementation Approach."""
        self.doc.add_heading("8. Implementation Approach", level=1)
        self.doc.add_paragraph(
            "[ Implementation methodology, phases, milestones, and resource plan "
            "to be detailed by the delivery team. ]"
        )
        self.doc.add_paragraph("")

    def _add_project_plan(self, ctx: dict):
        """Section 9: Project Plan."""
        self.doc.add_heading("9. Project Plan", level=1)

        schedule = ctx.get("schedule", [])
        if schedule:
            self.doc.add_heading("Key Milestones", level=2)
            table = self.doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Event"
            hdr[1].text = "Date"
            hdr[2].text = "Notes"

            for event in schedule:
                row = table.add_row().cells
                row[0].text = event.get("event_name", "N/A")
                row[1].text = str(event.get("date") or event.get("event_date", "TBD"))
                row[2].text = event.get("notes", "")
        else:
            self.doc.add_paragraph("[ Project timeline to be developed based on RFP schedule. ]")

        self.doc.add_paragraph("")

    def _add_pricing(self, ctx: dict):
        """Section 10: Pricing."""
        self.doc.add_heading("10. Pricing", level=1)

        pricing = ctx.get("pricing", [])
        if pricing:
            table = self.doc.add_table(rows=1, cols=5)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Category"
            hdr[1].text = "Line Item"
            hdr[2].text = "Unit Cost"
            hdr[3].text = "Quantity"
            hdr[4].text = "Total"

            for item in pricing:
                row = table.add_row().cells
                row[0].text = (item.get("category") or "").replace("_", " ").title()
                row[1].text = item.get("line_item", "N/A")
                row[2].text = f"${item.get('unit_cost', 0):,.2f}" if item.get("unit_cost") else "TBD"
                row[3].text = str(item.get("quantity", "")) if item.get("quantity") else ""
                row[4].text = f"${item.get('total', 0):,.2f}" if item.get("total") else "TBD"
        else:
            self.doc.add_paragraph(
                "[ Pricing details to be completed by the finance team based on "
                "the pricing template provided in the RFP. ]"
            )

        self.doc.add_paragraph("")

    def _add_assumptions(self, ctx: dict):
        """Section 11: Assumptions."""
        self.doc.add_heading("11. Assumptions", level=1)
        assumptions = ctx.get("assumptions", [
            "Client will provide timely access to stakeholders for requirements clarification.",
            "Existing infrastructure meets minimum system requirements.",
            "Client will designate a project sponsor and project manager.",
            "Standard business hours for implementation support unless otherwise agreed.",
            "Data migration scope is limited to data formats specified in the RFP.",
        ])
        for i, assumption in enumerate(assumptions, 1):
            self.doc.add_paragraph(f"{i}. {assumption}")
        self.doc.add_paragraph("")

    def _add_risks_mitigation(self, ctx: dict):
        """Section 12: Risks & Mitigation."""
        self.doc.add_heading("12. Risks & Mitigation", level=1)

        risks = ctx.get("risks", [
            {"risk": "Scope creep during implementation", "mitigation": "Clear change management process with formal change requests"},
            {"risk": "Data migration complexity", "mitigation": "Dedicated data migration phase with validation checkpoints"},
            {"risk": "User adoption challenges", "mitigation": "Comprehensive training program and change management support"},
            {"risk": "Integration complexity with existing systems", "mitigation": "Early integration testing and dedicated integration team"},
        ])

        table = self.doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "#"
        hdr[1].text = "Risk"
        hdr[2].text = "Mitigation"

        for i, risk in enumerate(risks, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = risk.get("risk", "")
            row[2].text = risk.get("mitigation", "")

        self.doc.add_paragraph("")

    def _add_legal_compliance(self, ctx: dict):
        """Section 13: Legal & Compliance Statements."""
        self.doc.add_heading("13. Legal & Compliance Statements", level=1)
        self.doc.add_paragraph(
            "[ Legal and compliance statements to be reviewed and finalized by the legal team. "
            "Include data protection commitments, regulatory compliance, "
            "confidentiality agreements, and standard legal terms. ]"
        )
        self.doc.add_paragraph("")

    def _add_footer(self, ctx: dict):
        """Add document footer with version info."""
        section = self.doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run(
            f"CONFIDENTIAL | {ctx.get('company', {}).get('name', 'Company')} | "
            f"Version 1.0 | {datetime.now().strftime('%Y-%m-%d')}"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def generate_word_document(context: dict[str, Any]) -> bytes:
    """Convenience function to generate a Word document."""
    generator = WordGenerator()
    return generator.generate(context)
