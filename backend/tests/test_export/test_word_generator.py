import pytest
import io
from docx import Document
from app.export.word_generator import generate_word_document


class TestWordGenerator:
    def _make_context(self, **overrides):
        context = {
            "project": {
                "name": "Test RFP Project",
                "description": "A test project",
                "client_name": "Test Client",
                "industry": "Technology",
                "status": "in_progress",
            },
            "company": {
                "name": "Test Company",
                "description": "We are a test company",
            },
            "requirements": [
                {
                    "id": "req-1",
                    "req_number": "FR-001",
                    "title": "User Authentication",
                    "description": "System shall support SSO authentication",
                    "type": "functional",
                    "category": "security",
                    "is_mandatory": True,
                    "priority": "high",
                },
                {
                    "id": "req-2",
                    "req_number": "NFR-001",
                    "title": "99.9% Uptime",
                    "description": "System shall maintain 99.9% availability",
                    "type": "non_functional",
                    "category": "availability",
                    "is_mandatory": True,
                    "priority": "high",
                },
            ],
            "responses": [
                {
                    "requirement_id": "req-1",
                    "compliance_status": "fully_compliant",
                    "response_text": "Our platform natively supports SSO via SAML 2.0 and OAuth 2.0.",
                    "confidence_score": 0.92,
                    "is_reviewed": True,
                },
                {
                    "requirement_id": "req-2",
                    "compliance_status": "fully_compliant",
                    "response_text": "We guarantee 99.95% uptime SLA with redundant infrastructure.",
                    "confidence_score": 0.88,
                    "is_reviewed": False,
                },
            ],
            "schedule": [
                {
                    "event_name": "Submission Deadline",
                    "event_type": "submission_deadline",
                    "date": "2026-03-15",
                    "notes": "",
                },
            ],
            "pricing": [],
            "compliance_scores": {
                "overall_score": 100.0,
                "functional_score": 100.0,
                "non_functional_score": 100.0,
                "total_requirements": 2,
                "total_responses": 2,
                "status_breakdown": {"fully_compliant": 2},
            },
        }
        context.update(overrides)
        return context

    def test_generates_valid_docx(self):
        context = self._make_context()
        doc_bytes = generate_word_document(context)
        assert doc_bytes is not None
        assert len(doc_bytes) > 0

        # Verify it's a valid DOCX
        doc = Document(io.BytesIO(doc_bytes))
        assert len(doc.paragraphs) > 0

    def test_contains_project_name(self):
        context = self._make_context()
        doc_bytes = generate_word_document(context)
        doc = Document(io.BytesIO(doc_bytes))

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Test RFP Project" in all_text

    def test_contains_all_sections(self):
        context = self._make_context()
        doc_bytes = generate_word_document(context)
        doc = Document(io.BytesIO(doc_bytes))

        all_text = "\n".join(p.text for p in doc.paragraphs)
        expected_sections = [
            "Executive Summary",
            "About the Company",
            "Understanding of Requirements",
            "Proposed Solution Overview",
            "Functional Compliance Matrix",
            "Non-Functional Compliance Matrix",
            "Architecture Overview",
            "Implementation Approach",
            "Project Plan",
            "Pricing",
            "Assumptions",
            "Risks & Mitigation",
            "Legal & Compliance",
        ]
        for section in expected_sections:
            assert section in all_text, f"Missing section: {section}"

    def test_contains_tables(self):
        context = self._make_context()
        doc_bytes = generate_word_document(context)
        doc = Document(io.BytesIO(doc_bytes))
        assert len(doc.tables) > 0

    def test_empty_requirements(self):
        context = self._make_context(requirements=[], responses=[])
        doc_bytes = generate_word_document(context)
        assert doc_bytes is not None
        assert len(doc_bytes) > 0
